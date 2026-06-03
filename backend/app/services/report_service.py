"""Report generation — DOCX incident reports + CSV export per spec section 13."""
import os
import csv
import io
from datetime import datetime
from typing import List, Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

from app.core.config import get_settings

settings = get_settings()


def generate_incident_docx(event: dict) -> bytes:
    """Generate DOCX incident report — spec section 13."""
    if not _DOCX_AVAILABLE:
        return _placeholder_docx(event)

    doc = Document()

    # Title
    title = doc.add_heading("INCIDENT REPORT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("SurveillanceIQ — AI Surveillance Intelligence Platform", 2)
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}")
    doc.add_paragraph()

    # Event summary table
    doc.add_heading("Incident Details", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"

    severity_color = {
        "L1": RGBColor(0x22, 0xC5, 0x5E),
        "L2": RGBColor(0xF5, 0x9E, 0x0B),
        "L2+": RGBColor(0xF9, 0x73, 0x16),
        "L3": RGBColor(0xEF, 0x44, 0x44),
    }

    rows_data = [
        ("Event ID", event.get("event_id", "N/A")),
        ("Event Type", event.get("event_type", "N/A").replace("_", " ").title()),
        ("Severity", event.get("severity", "N/A")),
        ("Threat Score", f"{event.get('threat_score', 0.0):.4f}"),
        ("Timestamp", str(event.get("timestamp", "N/A"))),
        ("Zone", event.get("zone_id", "N/A")),
        ("Video ID", event.get("video_id", "N/A")),
        ("Acknowledged", str(event.get("acknowledged", False))),
    ]
    for field, val in rows_data:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = str(val)

    doc.add_paragraph()

    # GenAI Summary
    genai = event.get("genai_data") or {}
    if genai:
        doc.add_heading("AI Incident Analysis", 1)
        doc.add_heading("Incident Summary", 2)
        doc.add_paragraph(genai.get("incident_summary", "N/A"))
        doc.add_heading("Classification Reasoning", 2)
        doc.add_paragraph(genai.get("classification_reasoning", "N/A"))
        doc.add_heading("Recommended Action", 2)
        para = doc.add_paragraph(genai.get("recommended_action", "N/A"))
        para.runs[0].bold = True
        doc.add_heading("Confidence Note", 2)
        doc.add_paragraph(genai.get("confidence_note", "N/A"))
        doc.add_paragraph()

    # Persons involved
    persons = event.get("persons_involved") or []
    if persons:
        doc.add_heading("Persons Involved", 1)
        ptable = doc.add_table(rows=1, cols=5)
        ptable.style = "Table Grid"
        hdrs = ptable.rows[0].cells
        for i, h in enumerate(["Track ID", "Person ID", "Name", "Confidence", "Duration (s)"]):
            hdrs[i].text = h
        for p in persons:
            row = ptable.add_row().cells
            row[0].text = str(p.get("track_id", ""))
            row[1].text = str(p.get("person_id", ""))
            row[2].text = str(p.get("display_name", ""))
            row[3].text = f"{p.get('face_confidence', 0.0):.3f}"
            row[4].text = f"{p.get('duration_in_zone_seconds', 0):.0f}"
        doc.add_paragraph()

    # Threat features
    features = event.get("threat_features") or {}
    if features:
        doc.add_heading("Threat Feature Vector", 1)
        ftable = doc.add_table(rows=1, cols=2)
        ftable.style = "Table Grid"
        fhdr = ftable.rows[0].cells
        fhdr[0].text = "Feature"
        fhdr[1].text = "Value"
        for k, v in features.items():
            frow = ftable.add_row().cells
            frow[0].text = k.replace("_", " ").title()
            frow[1].text = str(round(v, 4) if isinstance(v, float) else v)

    doc.add_paragraph()
    doc.add_paragraph("— End of Incident Report —").alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _placeholder_docx(event: dict) -> bytes:
    """Plain text fallback when python-docx unavailable."""
    content = f"""INCIDENT REPORT
===============
Event ID: {event.get('event_id')}
Type: {event.get('event_type')}
Severity: {event.get('severity')}
Threat Score: {event.get('threat_score')}
Timestamp: {event.get('timestamp')}
"""
    return content.encode("utf-8")


def generate_events_csv(events: List[dict]) -> str:
    """Generate CSV export — spec section 13."""
    output = io.StringIO()
    fieldnames = [
        "event_id", "video_id", "event_type", "severity", "threat_score",
        "timestamp", "zone_id", "persons_involved_count",
        "acknowledged", "acknowledged_by", "acknowledged_at",
        "incident_summary", "recommended_action",
    ]
    writer = csv.DictWriter(output, fieldnames=fieldnames, extrasaction="ignore")
    writer.writeheader()
    for e in events:
        genai = e.get("genai_data") or {}
        persons = e.get("persons_involved") or []
        writer.writerow({
            "event_id": e.get("event_id", e.get("id", "")),
            "video_id": e.get("video_id", ""),
            "event_type": e.get("event_type", ""),
            "severity": e.get("severity", ""),
            "threat_score": e.get("threat_score", 0.0),
            "timestamp": e.get("timestamp", ""),
            "zone_id": e.get("zone_id", ""),
            "persons_involved_count": len(persons),
            "acknowledged": e.get("acknowledged", False),
            "acknowledged_by": e.get("acknowledged_by", ""),
            "acknowledged_at": e.get("acknowledged_at", ""),
            "incident_summary": genai.get("incident_summary", ""),
            "recommended_action": genai.get("recommended_action", ""),
        })
    return output.getvalue()
